from docx import Document
import fitz  # PyMuPDF
import os
import pytesseract
from PIL import Image
import io
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Optional: Configure path to tesseract.exe (Windows only)
# Uncomment and adjust path as needed for your system
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

def extract_text_from_pdf(path):
    """
    Extract text from PDF using both direct text extraction and OCR
    """
    try:
        text = ""
        ocr_text = ""
        
        doc = fitz.open(path)
        logger.info(f"Processing PDF with {len(doc)} pages")
        
        for page_num, page in enumerate(doc):
            try:
                # Extract visible text
                page_text = page.get_text()
                text += page_text
                
                # If very little text found, try OCR
                if len(page_text.strip()) < 50:
                    try:
                        # Render page as image
                        pix = page.get_pixmap(dpi=300)
                        img = Image.open(io.BytesIO(pix.tobytes("png")))
                        
                        # Run OCR on image
                        ocr_page_text = pytesseract.image_to_string(img, lang='eng')
                        ocr_text += f"\n[Page {page_num + 1} OCR]\n" + ocr_page_text
                        
                    except Exception as ocr_error:
                        logger.warning(f"OCR failed for page {page_num + 1}: {str(ocr_error)}")
                        
            except Exception as page_error:
                logger.error(f"Error processing page {page_num + 1}: {str(page_error)}")
                continue
        
        doc.close()
        
        # Combine text and OCR results
        final_text = text
        if ocr_text.strip():
            final_text += "\n\n[OCR-Extracted Text]\n" + ocr_text
            
        return final_text.strip() if final_text.strip() else "No text could be extracted from this PDF."
        
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        raise Exception(f"Failed to extract text from PDF: {str(e)}")

def extract_text_from_docx(path):
    """
    Extract text from DOCX file
    """
    try:
        doc = Document(path)
        paragraphs = []
        
        for para in doc.paragraphs:
            if para.text.strip():  # Only add non-empty paragraphs
                paragraphs.append(para.text.strip())
        
        # Also extract text from tables
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    tables_text.append(" | ".join(row_text))
        
        # Combine paragraphs and tables
        all_text = []
        if paragraphs:
            all_text.extend(paragraphs)
        if tables_text:
            all_text.append("\n[Tables Content]\n")
            all_text.extend(tables_text)
        
        result = "\n".join(all_text)
        return result if result.strip() else "No text could be extracted from this DOCX file."
        
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")

def extract_text_from_txt(path):
    """
    Extract text from TXT file with multiple encoding attempts
    """
    encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
    
    for encoding in encodings:
        try:
            with open(path, "r", encoding=encoding) as f:
                content = f.read()
                if content.strip():
                    return content
                else:
                    return "The text file appears to be empty."
        except UnicodeDecodeError:
            continue
        except Exception as e:
            logger.error(f"Error reading TXT file with {encoding}: {str(e)}")
            continue
    
    raise Exception("Failed to read text file with any supported encoding")

def extract_text(path):
    """
    Main function to extract text from various file formats
    """
    if not os.path.exists(path):
        raise FileNotFoundError(f"File not found: {path}")
    
    if os.path.getsize(path) == 0:
        raise ValueError("File is empty")
    
    ext = os.path.splitext(path)[-1].lower()
    
    logger.info(f"Extracting text from {ext} file: {os.path.basename(path)}")
    
    if ext == ".pdf":
        return extract_text_from_pdf(path)
    elif ext == ".docx":
        return extract_text_from_docx(path)
    elif ext == ".txt":
        return extract_text_from_txt(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported formats: PDF, DOCX, TXT")

# Test function
def test_extraction(file_path):
    """
    Test function to verify text extraction works
    """
    try:
        text = extract_text(file_path)
        print(f"Successfully extracted {len(text)} characters")
        print(f"Preview: {text[:200]}...")
        return True
    except Exception as e:
        print(f"Extraction failed: {str(e)}")
        return False

if __name__ == "__main__":
    # Example usage
    print("Text Extractor Module")
    print("Supported formats: PDF, DOCX, TXT")